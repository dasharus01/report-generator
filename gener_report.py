from docxtpl import DocxTemplate
import io
# для года
import datetime
# для получения данных из консоли
import sys
from docx import Document
from docx.shared import Inches
import argparse
import json
import pickle
parser = argparse.ArgumentParser()

parser.add_argument('-o', '--org', help="Info about organisation")
parser.add_argument('-s', '--stu', help="Info about student")
parser.add_argument('-r', '--rep' , help = "Info about report")
parser.add_argument('-t', '--tea' , help = "Info about teacher")
parser.add_argument('--struc' , help = "Info about paragraph")
# год текущий
now = datetime.datetime.now()
year = now.year 
def pr (d):
    doc = DocxTemplate('template.docx')
    founder = "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ"
    nameuni = "федеральное государственное автономное образовательное учреждение высшего образования «САНКТ-ПЕТЕРБУРГСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ АЭРОКОСМИЧЕСКОГО ПРИБОРОСТРОЕНИЯ»"
    if not d['organisation'].get('department', ""):
        d['organisation']['department'] = ""
    if not d['report'].get('task_type', ""):
        d['report']['task_type'] = ""
        
    if not d['report'].get('task_name', ""):
        d['report']['task_name'] = ""
        
    if  not d['report'].get('subject_name', ""):
        d['report']['subject_name'] = ""
    if d['organisation'].get('name', "") and d['organisation'].get('founder', ""):
        nameuni = d['organisation']['name']
        founder = d['organisation']['founder']
    # поднимаем регистор
    d['organisation']['department'] = d['organisation']['department'].upper()
    d['report']['task_type'] = d['report']['task_type'].upper()
    d['report']['task_name'] = d['report']['task_name'].upper()
    d['report']['subject_name'] = d['report']['subject_name'].upper()
    #преобразуем имя студента
    if not d['student'].get('surname', ""):
        d['student']['surname'] = " "
    if not d['student'].get('name', ""):
        d['student']['name'] = " "
    
    if not d['student'].get('group', ""):
        d['student']['group'] = " "
    
    if not d['teacher'].get('surname', "") :
        d['teacher']['surname'] = " "
        
    if not d['teacher'].get('name', ""):
        d['teacher']['name'] = " "
        
    if not d['teacher'].get('status', ""):
        d['teacher']['status'] = " "
    
    if d['student'].get('patronymic', ""):
        names = d['student']['surname'] + " " + d['student']['name'][0] + "." + d['student']['patronymic'][0] + "."
    else:
        names = d['student']['surname'] + " " + d['student']['name'][0] + "."
  
    
    # преобразуем имя преподавателя
    if d['teacher'].get('patronymic', ""):
        teachern = d['teacher']['surname'] + " " + d['teacher']['name'][0] + "." + d['teacher']['patronymic'][0] + "."
    else:
        teachern = d['teacher']['surname'] + " " + d['teacher']['name'][0] + "."
    # наличие студ.билета
    studbilet = ""
    identity_card = ""
    if d['student'].get('identity_card', ""):
        studbilet = "Студенческий билет №"
        identity_card = d['student']['identity_card']
    # наличие факультета
    faculty = ""
    if d['organisation'].get('faculty', ""):
        faculty = d['organisation']['faculty'].upper()
    # заполнение атоматической части
    if d['report']['task_type'].find("КОНТРОЛЬНАЯ") != -1:
        type2 = ""
        point = "ОЦЕНКА"
        lecturer = "ПРЕПОДАВАТЕЛЬ"
        how = "по дисциплине"
        
    elif d['report']['task_type'].find("ПРОЕКТУ") != -1:
        type2 = "КУРСОВОЙ ПРОЕКТ"
        point = "ЗАЩИЩЕН С ОЦЕНКОЙ"
        lecturer = "РУКОВОДИТЕЛЬ"
        how = "по дисциплине"
       
    elif d['report']['task_type'].find("КУРСОВОЙ РАБОТЕ ") != -1:
        type2 = "КУРСОВОЙ РАБОТА"
        point = "ЗАЩИЩЕНА С ОЦЕНКОЙ"
        lecturer = "РУКОВОДИТЕЛЬ"
        how = "по дисциплине"
        
    elif d['report']['task_type'].find("ЛАБОР") != -1:
        type2 = "ОТЧЕТ"
        point = "ЗАЩИЩЕН С ОЦЕНКОЙ"
        lecturer = "ПРЕПОДАВАТЕЛЬ"
        how = "по курсу"
        
        
    elif d['report']['task_type'].find("РЕФЕРАТ") != -1:
        type2 = ""
        point = "ОЦЕНКА РЕФЕРАТА"
        lecturer = "РУКОВОДИТЕЛЬ"
        how = "по дисциплине"
    # stri = d["report_structure"][0]
    
    
    # заполнение шаблона
    context = {'founder' : founder,
               'name' : nameuni,
               'faculty' : faculty,
               'department' : d['organisation']['department'],
               'namestudent' : names,
               'group' : d['student']['group'],
               'ids' : identity_card,
               'studbilet' : studbilet,
               'typework' : d['report']['task_type'],
               'namework' : d['report']['task_name'],
               'subject' : d['report']['subject_name'],
               'status' : d['teacher']['status'],
               'teachername' : teachern,
               'year' : year,
               'type2' : type2, 'point' : point, 'lecturer' : lecturer, 'how' : how
              }
    doc.render(context)
    i = 0;
    lend = len(d["report_structure"])
    while lend != 0:
        doc.add_heading(d["report_structure"][i], level=1)
        doc.add_paragraph('')
        i = i +1
        lend = lend - 1 
    ######doc.save('result2.docx')
    #document = Document()
    
    #document = Document()
    #doc1 = open('report.docx', 'rb').read()
   
    ##############
    # Create in-memory buffer
    file_stream = io.BytesIO()
    # Save the .docx to the buffer
    doc.save(file_stream)
    # Reset the buffer's file-pointer to the beginning of the file
    file_stream.seek(0)
    print(file_stream)
    
    return file_stream
    
def main():
    d = {}
    args = parser.parse_args()
    d["organisation"] = json.loads(args.org)
    d["student"] = json.loads(args.stu)
    d["report"] = json.loads(args.rep)
    d["teacher"] = json.loads(args.tea)
    d["report_structure"] = json.loads(args.struc)
   
  
    with open("result.docx", "wb") as f:
        for i in pr(d):
            f.write(i)
    #print(fp)
if __name__ == "__main__":
    main()
    
