import gener_report
from docxtpl import DocxTemplate
import io
# для года
import datetime
# для получения данных из консоли
import sys
from docx import Document
from docx.shared import Inches
import filecmp
import pathlib
import hashlib
import codecs
def main(d):
    
    #"status": "Ректор, д.т.н., проф."
    # открытие файла с титульником
    with open("result.docx", "wb") as f:
            for i in gener_report.pr(d):
                f.write(i)
    # файл с эталоном
    doc = Document("result2.docx")
    # файл с титульником
    doc1 = Document("result.docx")
    # выгрузка таблиц
    table1 = doc.tables
    table2 = doc1.tables
    # выгрузка строк
    all_paras = doc.paragraphs
    all_paras1 = doc1.paragraphs
    # определяем самый маленький файл
    k = 0;
    if len(all_paras) < len(all_paras1):
        k = len(all_paras)
    else:
        k = len(all_paras1)
    #пробег по строкам
    for i in range(0, k):
        if all_paras[i].text != all_paras1[i].text:
            if i <= 3:
                sys.exit(1)
            else:
                sys.exit(2)
    if len(all_paras) != len(all_paras1):
        sys.exit(3)
        


    # пробег по таблицам
    row = table2[0].rows[0]
    cell = row.cells
    row1 = table1[0].rows[0]
    cell1 = row1.cells
    #статус
    if cell[0].text != cell1[0].text:
        sys.exit(4)
    # ФИО препода  
    if cell[4].text != cell1[4].text:
        sys.exit(5)

    row = table2[1].rows[0]
    cell = row.cells
    row1 = table1[1].rows[0]
    cell1 = row1.cells
    # тип работы 
    if cell[0].text != cell1[0].text:
        sys.exit(6)

    row = table2[1].rows[1]
    cell = row.cells
    row1 = table1[1].rows[1]
    cell1 = row1.cells
    # название работы 
    if cell[0].text != cell1[0].text:
        sys.exit(7)
    row = table2[1].rows[2]
    cell = row.cells
    row1 = table1[1].rows[2]
    cell1 = row1.cells
    # название предмета  
    if cell[0].text != cell1[0].text:
        sys.exit(8)


    row = table2[2].rows[0]
    cell = row.cells
    row1 = table1[2].rows[0]
    cell1 = row1.cells
    # группа
    if cell[1].text != cell1[1].text:
        sys.exit(9)
    # ФИО студента   
    if cell[5].text != cell1[5].text:
        sys.exit(10)   
